import csv
import hashlib
import json
import shutil
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "车辆违停YOLO11_完整标注数据集"
PROJECT = ROOT / "parking_violation_yolo11_project"
ASSETS = PROJECT / "paper_assets"
OUT = ROOT / "Campus_Illegal_Parking_Dual_View_Dataset_ScientificData_Package"

DATASET_TITLE = "Campus Illegal Parking Dual-view Vehicle Detection Dataset"
CLASS_NAME = "illegal_parked_vehicle"
CLASS_ID = 0
LICENSE_TEXT = """Dataset license recommendation: CC BY-NC 4.0 for non-commercial academic use.

Before public release, the dataset owner should confirm the final license with the target repository and institution. The current package is prepared as a publication-ready release candidate, but it should not be uploaded until the privacy review checklist is signed off.
"""


ACQUISITION = {
    "collection_region": "Changqing University Town, Jinan, Shandong, China",
    "specific_locations_for_manuscript": [
        "Roads around Shandong Jiaotong University",
        "Roads around Shandong Management University",
        "Roads around Shandong Normal University",
        "Roads around Shandong University of Art and Design",
        "Roads around Shandong Women's University",
        "Roads around Qilu University of Technology",
        "Roads around Shandong University of Traditional Chinese Medicine",
    ],
    "roads_recorded_in_readme": [
        "Haitang Road",
        "Dingxiang Road",
        "Daxue Road",
        "Wate Road",
        "Ziwei Road",
        "Road between Jiguang Expressway and Shandong Management University",
    ],
    "collection_dates": ["2026-03-21", "2026-03-23", "2026-03-28", "2026-04-18", "2026-04-19", "2026-04-21"],
    "weather": "daytime, sunny or cloudy",
    "front_view_device": "Huawei nova 12 Pro smartphone built-in camera",
    "top_view_devices": [
        "DJI Flip 2 UAV video frames",
        "selfie-stick overhead photographs",
    ],
    "top_view_height": "approximately 2-3 m above the parked vehicle",
    "top_view_sampling": "high-quality frames with large vehicle area were selected from street-level UAV video recordings; some top-view images were directly photographed using a selfie stick",
    "privacy_processing": "selfie-stick shadows and visible photographer information were blurred when present; public-release files use de-identified filenames and EXIF metadata removal",
}


def sha256(path):
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def infer_view(stem):
    s = stem.lower()
    if "front" in s:
        return "front"
    if "top" in s:
        return "top"
    return "unknown"


def infer_group_id(stem):
    digits = []
    for ch in stem:
        if ch.isdigit():
            digits.append(ch)
        else:
            break
    return "".join(digits) or stem


def read_yolo_label(path):
    rows = []
    if not path.exists():
        return rows
    for line in path.read_text(encoding="utf-8").splitlines():
        if not line.strip():
            continue
        parts = line.split()
        if len(parts) != 5:
            rows.append({"bad": True, "raw": line})
            continue
        cls, x, y, w, h = parts
        rows.append({"class_id": int(cls), "x": float(x), "y": float(y), "w": float(w), "h": float(h)})
    return rows


def save_without_exif(src, dst):
    dst.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(src) as im:
        im = im.convert("RGB") if im.mode not in ("RGB", "L") else im.copy()
        if dst.suffix.lower() in [".jpg", ".jpeg"]:
            im.save(dst, quality=95, optimize=True)
        else:
            im.save(dst)


def prepare_dirs():
    if OUT.exists():
        shutil.rmtree(OUT)
    for p in [
        OUT / "data" / "images",
        OUT / "data" / "labels_yolo",
        OUT / "data" / "annotations_coco",
        OUT / "data" / "splits",
        OUT / "metadata",
        OUT / "code",
        OUT / "docs",
        OUT / "paper" / "figures",
        OUT / "paper" / "tables",
    ]:
        p.mkdir(parents=True, exist_ok=True)


def copy_scripts():
    for name in [
        "train_yolo11.py",
        "evaluate_yolo11.py",
        "make_final_paper_assets.py",
        "build_scientific_data_release.py",
    ]:
        src = PROJECT / name
        if src.exists():
            shutil.copy2(src, OUT / "code" / name)
    for fig in [
        "fig_split_distribution.png",
        "fig_view_distribution.png",
        "fig_ablation_map95.png",
        "fig_imgsz_ablation.png",
        "fig_model_scale_comparison.png",
        "fig_view_map95.png",
        "fig_prediction_examples.jpg",
        "results.png",
        "BoxPR_curve.png",
        "confusion_matrix.png",
    ]:
        src = ASSETS / fig
        if src.exists():
            shutil.copy2(src, OUT / "paper" / "figures" / fig)
    for tbl in ["table_dataset.csv", "table_results.csv", "table_imgsz_ablation.csv", "table_model_scale.csv"]:
        src = ASSETS / tbl
        if src.exists():
            shutil.copy2(src, OUT / "paper" / "tables" / tbl)


def build_release_data():
    images_meta = []
    label_meta = []
    qc = {
        "image_count": 0,
        "label_count": 0,
        "missing_labels": [],
        "empty_labels": [],
        "bad_label_lines": [],
        "bbox_out_of_range": [],
        "exif_removed_images": 0,
    }

    for split in ["train", "val", "test"]:
        img_dir = SOURCE / "images" / split
        label_dir = SOURCE / "labels" / split
        split_names = []
        for img_path in sorted([p for p in img_dir.iterdir() if p.suffix.lower() in [".jpg", ".jpeg", ".png"]]):
            rel_img = Path("data") / "images" / split / img_path.name
            dst_img = OUT / rel_img
            before_has_exif = False
            with Image.open(img_path) as im:
                width, height = im.size
                before_has_exif = bool(getattr(im, "getexif", lambda: {})())
            save_without_exif(img_path, dst_img)
            if before_has_exif:
                qc["exif_removed_images"] += 1

            label_path = label_dir / f"{img_path.stem}.txt"
            rel_label = Path("data") / "labels_yolo" / split / label_path.name
            dst_label = OUT / rel_label
            dst_label.parent.mkdir(parents=True, exist_ok=True)
            labels = read_yolo_label(label_path)
            if not label_path.exists():
                qc["missing_labels"].append(str(rel_img))
                dst_label.write_text("", encoding="utf-8")
            else:
                shutil.copy2(label_path, dst_label)
                qc["label_count"] += 1
            if not labels:
                qc["empty_labels"].append(str(rel_label))
            for idx, row in enumerate(labels):
                if row.get("bad"):
                    qc["bad_label_lines"].append({"file": str(rel_label), "line": row["raw"]})
                    continue
                vals = [row["x"], row["y"], row["w"], row["h"]]
                if any(v < 0 or v > 1 for v in vals) or row["w"] <= 0 or row["h"] <= 0:
                    qc["bbox_out_of_range"].append({"file": str(rel_label), "line_index": idx, "values": vals})
                label_meta.append({
                    "image_file": str(rel_img).replace("\\", "/"),
                    "label_file": str(rel_label).replace("\\", "/"),
                    "class_id": row.get("class_id"),
                    "class_name": CLASS_NAME,
                    "x_center": row.get("x"),
                    "y_center": row.get("y"),
                    "width": row.get("w"),
                    "height": row.get("h"),
                })

            split_names.append(str(rel_img).replace("\\", "/"))
            images_meta.append({
                "file_name": str(rel_img).replace("\\", "/"),
                "split": split,
                "group_id": infer_group_id(img_path.stem),
                "view": infer_view(img_path.stem),
                "width_px": width,
                "height_px": height,
                "extension": img_path.suffix.lower(),
                "sha256": sha256(dst_img),
                "label_file": str(rel_label).replace("\\", "/"),
                "object_count": len([x for x in labels if not x.get("bad")]),
            })
            qc["image_count"] += 1
        (OUT / "data" / "splits" / f"{split}.txt").write_text("\n".join(split_names) + "\n", encoding="utf-8")

    return images_meta, label_meta, qc


def build_coco(images_meta, label_meta):
    labels_by_image = defaultdict(list)
    for row in label_meta:
        labels_by_image[row["image_file"]].append(row)

    for split in ["train", "val", "test"]:
        images = []
        annotations = []
        img_id_by_file = {}
        ann_id = 1
        split_images = [r for r in images_meta if r["split"] == split]
        for idx, row in enumerate(split_images, 1):
            img_id_by_file[row["file_name"]] = idx
            images.append({
                "id": idx,
                "file_name": row["file_name"],
                "width": row["width_px"],
                "height": row["height_px"],
                "split": split,
                "view": row["view"],
                "group_id": row["group_id"],
            })
        for row in split_images:
            for lab in labels_by_image[row["file_name"]]:
                x = (lab["x_center"] - lab["width"] / 2) * row["width_px"]
                y = (lab["y_center"] - lab["height"] / 2) * row["height_px"]
                w = lab["width"] * row["width_px"]
                h = lab["height"] * row["height_px"]
                annotations.append({
                    "id": ann_id,
                    "image_id": img_id_by_file[row["file_name"]],
                    "category_id": CLASS_ID,
                    "bbox": [round(x, 3), round(y, 3), round(w, 3), round(h, 3)],
                    "area": round(w * h, 3),
                    "iscrowd": 0,
                    "segmentation": [],
                })
                ann_id += 1
        coco = {
            "info": {
                "description": DATASET_TITLE,
                "version": "1.0-release-candidate",
                "year": 2026,
                "date_created": date.today().isoformat(),
            },
            "licenses": [{"id": 1, "name": "License to be confirmed before public release"}],
            "images": images,
            "annotations": annotations,
            "categories": [{"id": CLASS_ID, "name": CLASS_NAME, "supercategory": "vehicle"}],
        }
        (OUT / "data" / "annotations_coco" / f"{split}.json").write_text(
            json.dumps(coco, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


def write_metadata(images_meta, label_meta, qc):
    headers = list(images_meta[0].keys())
    write_csv(OUT / "metadata" / "image_metadata.csv", headers, images_meta)
    write_csv(OUT / "metadata" / "annotation_metadata.csv", list(label_meta[0].keys()), label_meta)

    split_counts = Counter(r["split"] for r in images_meta)
    view_counts = Counter(r["view"] for r in images_meta)
    stats = [
        {"metric": "total_images", "value": len(images_meta)},
        {"metric": "total_annotations", "value": len(label_meta)},
        {"metric": "vehicle_groups", "value": len(set(r["group_id"] for r in images_meta))},
        {"metric": "train_images", "value": split_counts["train"]},
        {"metric": "val_images", "value": split_counts["val"]},
        {"metric": "test_images", "value": split_counts["test"]},
        {"metric": "front_images", "value": view_counts["front"]},
        {"metric": "top_images", "value": view_counts["top"]},
        {"metric": "unknown_view_images", "value": view_counts["unknown"]},
    ]
    write_csv(OUT / "metadata" / "dataset_statistics.csv", ["metric", "value"], stats)

    checksums = []
    for p in sorted((OUT / "data").rglob("*")):
        if p.is_file():
            checksums.append({"sha256": sha256(p), "path": str(p.relative_to(OUT)).replace("\\", "/")})
    write_csv(OUT / "metadata" / "sha256_checksums.csv", ["sha256", "path"], checksums)

    qc_summary = {
        **qc,
        "missing_label_count": len(qc["missing_labels"]),
        "empty_label_count": len(qc["empty_labels"]),
        "bad_label_line_count": len(qc["bad_label_lines"]),
        "bbox_out_of_range_count": len(qc["bbox_out_of_range"]),
        "status": "PASS" if not qc["missing_labels"] and not qc["bad_label_lines"] and not qc["bbox_out_of_range"] else "REVIEW_REQUIRED",
    }
    (OUT / "metadata" / "quality_control_report.json").write_text(
        json.dumps(qc_summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def write_csv(path, headers, rows):
    with path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=headers)
        writer.writeheader()
        for row in rows:
            writer.writerow(row)


def write_docs():
    readme = f"""# {DATASET_TITLE}

This release-candidate package contains a dual-view image dataset for suspicious illegal parking vehicle detection around university roads in Changqing University Town, Jinan, Shandong, China.

## Scope

The dataset supports vehicle localization in suspicious illegal-parking scenes. It should not be described as a complete law-enforcement-grade illegal/legal parking classification dataset because it currently does not include systematic legal-parking negative samples or pixel-level road-context labels.

## Data format

- YOLO labels: `data/labels_yolo/<split>/*.txt`
- COCO labels: `data/annotations_coco/train.json`, `val.json`, `test.json`
- Images: `data/images/<split>/*`
- Splits: `data/splits/*.txt`
- Metadata: `metadata/*.csv`

## Class definition

| id | class name | definition |
|---|---|---|
| 0 | `{CLASS_NAME}` | Vehicle target in a suspicious illegal-parking scene. The bounding box encloses the visible vehicle body. |

## Acquisition summary

- Dates: {', '.join(ACQUISITION['collection_dates'])}
- Weather: {ACQUISITION['weather']}
- Front-view images: Huawei nova 12 Pro smartphone built-in camera.
- Top-view images: DJI Flip 2 UAV video frames and selfie-stick overhead photographs.
- Top-view height: {ACQUISITION['top_view_height']}.
- Privacy: public-release filenames remove license-plate text; EXIF metadata is removed; visible photographer information and selfie-stick shadows should be blurred when present.

## Reuse notes

For a Scientific Data submission, upload this package to a recognized repository such as Zenodo, Figshare, Dryad, or another institution-approved data repository and obtain a DOI. Update the DOI fields in the manuscript before submission.
"""
    (OUT / "README.md").write_text(readme, encoding="utf-8")

    (OUT / "LICENSE_TO_BE_CONFIRMED.md").write_text(LICENSE_TEXT, encoding="utf-8")
    (OUT / "CITATION.cff").write_text(
        """cff-version: 1.2.0
title: "Campus Illegal Parking Dual-view Vehicle Detection Dataset"
message: "If you use this dataset, please cite the associated Data Descriptor and dataset DOI."
type: dataset
authors:
  - family-names: "To be completed"
    given-names: "To be completed"
year: 2026
repository-code: "To be completed"
doi: "To be completed after repository upload"
license: "To be confirmed"
""",
        encoding="utf-8",
    )
    (OUT / "dataset_description.json").write_text(
        json.dumps(
            {
                "title": DATASET_TITLE,
                "version": "1.0-release-candidate",
                "task": "single-class object detection",
                "class_names": [CLASS_NAME],
                "formats": ["YOLO", "COCO"],
                "acquisition": ACQUISITION,
                "limitations": [
                    "No systematic legal-parking negative samples are included.",
                    "No pixel-level sidewalk, curb, lane, or road-region annotations are included.",
                    "Before public release, a final human privacy review for license plates, faces, and photographer information is required.",
                ],
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    annotation = """# Annotation Guidelines

## Object class

`illegal_parked_vehicle`: a visible vehicle in a suspicious illegal-parking scene.

## Bounding-box rule

Draw the smallest axis-aligned bounding box that encloses the visible vehicle body. Do not include large background regions, road surface, curb, or sidewalk unless they are visually inseparable from the vehicle boundary.

## Occlusion and truncation

If part of the vehicle is occluded or outside the image, annotate the visible vehicle region only. For top-view images, include the full visible vehicle footprint rather than a small discriminative part.

## Current limitation

The label indicates suspicious illegal-parking vehicle localization, not a final legal/illegal decision. Full violation reasoning requires road-context annotations and legal-parking negative examples.
"""
    (OUT / "docs" / "annotation_guidelines.md").write_text(annotation, encoding="utf-8")

    privacy = """# Privacy Review Checklist

Before uploading the dataset to a public repository, complete this checklist manually.

- [ ] Filenames contain no license-plate numbers or personal identifiers.
- [ ] EXIF/GPS metadata has been removed from all released images.
- [ ] Clearly visible faces are blurred.
- [ ] Clearly readable license plates are blurred or otherwise de-identified if required by the target repository, institution, or applicable data policy.
- [ ] Visible photographer information, selfie-stick shadows that reveal identity, and reflections are checked and blurred when needed.
- [ ] Exact sensitive locations are generalized in the manuscript when appropriate.

This release package already removes EXIF metadata and uses de-identified filenames. Visual privacy review still requires human confirmation.
"""
    (OUT / "docs" / "privacy_review_checklist.md").write_text(privacy, encoding="utf-8")

    usage = """# Usage Notes

## YOLO training

Use the root of this release package as the dataset path. The `data.yaml` file points to `data/images/train`, `data/images/val`, and `data/images/test`.

## COCO usage

COCO annotations are provided in `data/annotations_coco/`. The `file_name` field stores paths relative to the package root.

## Recommended reporting

When reporting results, state that the task is suspicious illegal-parking vehicle detection. Do not claim full illegal/legal parking classification unless additional negative samples and road-context labels are added.
"""
    (OUT / "docs" / "usage_notes.md").write_text(usage, encoding="utf-8")

    data_yaml = """path: .
train: data/images/train
val: data/images/val
test: data/images/test
names:
  0: illegal_parked_vehicle
"""
    (OUT / "data.yaml").write_text(data_yaml, encoding="utf-8")


def add_p(doc, text="", bold=False, size=10.5, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Microsoft YaHei"
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(15 if level == 1 else 12)
    r.bold = True
    return p


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


def build_data_descriptor_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    add_p(doc, "A dual-view image dataset for suspicious illegal parking vehicle detection around university roads", True, 16, WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Data Descriptor manuscript draft for Scientific Data-style submission", False, 10, WD_ALIGN_PARAGRAPH.CENTER)

    add_h(doc, "Background & Summary")
    add_p(doc, "Illegal parking around university roads can affect pedestrian movement, non-motorized traffic, and roadside order. Publicly reusable image datasets that capture both vehicle appearance and overhead road-context cues in campus-surrounding scenes remain limited. This dataset provides front-view and top-view images of suspicious illegal-parking vehicles collected around roads in Changqing University Town, Jinan, Shandong, China.")
    add_p(doc, "The dataset contains 584 images organized into 277 vehicle groups. The released annotations support single-class vehicle detection with the class illegal_parked_vehicle. The dataset is intended for vehicle localization and road-context-oriented research. It should not be interpreted as a complete legal/illegal parking decision dataset because systematic legal-parking negative samples and pixel-level road-region labels are not yet included.")

    add_h(doc, "Methods")
    add_p(doc, "Images were collected during daytime under sunny or cloudy conditions. Collection dates recorded in the source readme are 21 March 2026, 23 March 2026, 28 March 2026, 18 April 2026, 19 April 2026, and 21 April 2026. Collection locations include roads around several universities in Changqing University Town, including Shandong Jiaotong University, Shandong Management University, Shandong Normal University, Shandong University of Art and Design, Shandong Women's University, Qilu University of Technology, and Shandong University of Traditional Chinese Medicine.")
    add_p(doc, "Front-view images were captured with the built-in camera of a Huawei nova 12 Pro smartphone. Top-view images were obtained either by extracting high-quality frames from DJI Flip 2 UAV videos or by direct overhead photography with a selfie stick. For overhead views, the camera was approximately 2-3 m above the parked vehicle. UAV video frames were selected when the image was sharp and the target vehicle occupied a relatively large area.")
    add_p(doc, "The public-release package uses de-identified filenames and removes EXIF metadata. Photographer information, visible selfie-stick shadows, and other identifiable traces should be blurred when present. A final human privacy review is required before public repository upload.")
    add_p(doc, "Annotations are provided in YOLO and COCO formats. The train/validation/test split was performed by vehicle group to reduce leakage from multiple views of the same vehicle.")

    add_h(doc, "Data Records")
    add_p(doc, "The release candidate package is organized as follows: data/images contains train, validation, and test images; data/labels_yolo contains YOLO-format labels; data/annotations_coco contains COCO JSON files; metadata contains image metadata, annotation metadata, dataset statistics, quality-control results, and SHA-256 checksums.")
    add_table(doc, ["Component", "Path", "Description"], [
        ["Images", "data/images/<split>/", "De-identified image files with EXIF metadata removed"],
        ["YOLO labels", "data/labels_yolo/<split>/", "Single-class normalized bounding-box labels"],
        ["COCO labels", "data/annotations_coco/*.json", "COCO-format annotations for broader reuse"],
        ["Splits", "data/splits/*.txt", "Relative image paths for train, validation, and test"],
        ["Metadata", "metadata/*.csv/json", "Dataset statistics, checksums, and quality-control records"],
    ])

    add_h(doc, "Technical Validation")
    add_p(doc, "The release package includes automated integrity checks for image readability, image-label pairing, label format, normalized coordinate range, object counts, split counts, view counts, and SHA-256 checksums. Baseline YOLO11 experiments were conducted as a technical validation of annotation usability.")
    add_table(doc, ["Experiment", "Precision", "Recall", "mAP@0.5", "mAP@0.5:0.95", "Inference ms/img"], [
        ["YOLO11n, 416, test", "0.9861", "0.9890", "0.9864", "0.8726", "17.6"],
        ["YOLO11n, 320, test", "0.9890", "0.9848", "0.9846", "0.8361", "7.1"],
        ["YOLO11n, 512, test", "0.9782", "0.9868", "0.9859", "0.8618", "18.4"],
        ["YOLO11s, 416, test", "0.9889", "0.9777", "0.9880", "0.8605", "31.2"],
    ])
    add_p(doc, "These experiments are provided to validate that the annotations can train a standard detector and to characterize baseline performance. They are not presented as a claim of complete illegal-parking law-enforcement recognition.")

    add_h(doc, "Usage Notes")
    add_p(doc, "The dataset is suitable for suspicious illegal-parking vehicle localization, view-specific detection analysis, small-object/overhead vehicle detection studies, and road-context-aware parking research. Users who require legal/illegal parking classification should add normal-parking negative samples and road-context labels such as sidewalks, curbs, lanes, and no-parking signs.")

    add_h(doc, "Code Availability")
    add_p(doc, "Code for dataset validation, YOLO/COCO organization, baseline training, evaluation, and manuscript asset generation is included in the code directory of the release package. Repository URL and version tag should be added after GitHub or institutional repository upload.")

    add_h(doc, "Data Availability")
    add_p(doc, "Before submission, upload the release package to a recognized repository such as Zenodo, Figshare, Dryad, or an institutional repository and obtain a DOI. Replace this paragraph with the final repository citation and DOI.")

    add_h(doc, "References")
    refs = [
        "Scientific Data. Submission Guidelines. https://www.nature.com/sdata/submission-guidelines",
        "Ultralytics. YOLO11 Models Documentation. https://docs.ultralytics.com/models/yolo11/",
        "Real-Time Illegal Parking Detection System Based on Deep Learning. arXiv:1710.02546. https://arxiv.org/abs/1710.02546",
    ]
    for i, ref in enumerate(refs, 1):
        add_p(doc, f"[{i}] {ref}")

    doc.save(OUT / "paper" / "Scientific_Data_Data_Descriptor_Draft.docx")


def main():
    prepare_dirs()
    images_meta, label_meta, qc = build_release_data()
    build_coco(images_meta, label_meta)
    write_metadata(images_meta, label_meta, qc)
    copy_scripts()
    write_docs()
    build_data_descriptor_docx()
    print(OUT)


if __name__ == "__main__":
    main()
